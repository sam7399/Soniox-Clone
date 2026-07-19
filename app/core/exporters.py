"""Transcript exporters: TXT, SRT, VTT, JSON, Markdown, HTML, DOCX,
XLSX, PDF. All exporters read from the database and honour options for
timestamps, speakers, summary and translation inclusion."""
from __future__ import annotations

import datetime as dt
import html as html_mod
import json
import logging
from dataclasses import dataclass, field
from pathlib import Path

from app.config import APP_NAME
from app.db.database import db_session
from app.db.models import Session, Summary, Translation

log = logging.getLogger(__name__)


@dataclass
class ExportOptions:
    include_timestamps: bool = True
    include_speakers: bool = True
    include_summary: bool = False
    include_translation: str = ""       # target language code or ""
    title: str = ""
    watermark: str = ""                 # e.g. "CONFIDENTIAL"
    footer: str = ""


@dataclass
class SessionData:
    name: str
    date: str
    duration_s: float
    languages: str
    provider: str
    segments: list[dict] = field(default_factory=list)
    translation: list[dict] = field(default_factory=list)
    summary: dict = field(default_factory=dict)


def load_session_data(session_id: int, opts: ExportOptions) -> SessionData:
    with db_session() as s:
        sess = s.get(Session, session_id)
        if sess is None:
            raise ValueError("session not found")
        segments = [{
            "start_s": seg.start_s, "end_s": seg.end_s,
            "speaker": (seg.speaker.display_name or seg.speaker.label)
            if seg.speaker else "",
            "language": seg.language, "text": seg.text,
        } for seg in sess.segments]
        data = SessionData(
            name=sess.name,
            date=sess.created_at.strftime("%Y-%m-%d %H:%M"),
            duration_s=sess.duration_s, languages=sess.language,
            provider=sess.provider, segments=segments)
        if opts.include_translation:
            tr = next((t for t in sess.translations
                       if t.target_language == opts.include_translation),
                      None)
            if tr:
                data.translation = list(tr.segments_json)
        if opts.include_summary and sess.summaries:
            data.summary = dict(sess.summaries[-1].content_json)
    return data


def fmt_ts(seconds: float, vtt: bool = False) -> str:
    ms = int(round((seconds - int(seconds)) * 1000))
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    sep = "." if vtt else ","
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


def _line(seg: dict, opts: ExportOptions) -> str:
    parts = []
    if opts.include_timestamps:
        parts.append(f"[{fmt_ts(seg['start_s'])[:8]}]")
    if opts.include_speakers and seg.get("speaker"):
        parts.append(f"{seg['speaker']}:")
    parts.append(seg["text"])
    return " ".join(parts)


# ------------------------------------------------------------ plain text

def export_txt(session_id: int, out: Path, opts: ExportOptions) -> Path:
    d = load_session_data(session_id, opts)
    lines = [d.name, f"{d.date}  |  {fmt_ts(d.duration_s)[:8]}  |  "
             f"{d.languages}", ""]
    lines += [_line(seg, opts) for seg in d.segments]
    if d.summary:
        lines += ["", "=" * 40, "AI-GENERATED SUMMARY", "=" * 40,
                  _summary_text(d.summary)]
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def export_srt(session_id: int, out: Path, opts: ExportOptions) -> Path:
    d = load_session_data(session_id, opts)
    src = d.translation if opts.include_translation and d.translation \
        else d.segments
    blocks = []
    for i, seg in enumerate(src, 1):
        text = seg["text"]
        if opts.include_speakers and seg.get("speaker"):
            text = f"{seg['speaker']}: {text}"
        blocks.append(f"{i}\n{fmt_ts(seg['start_s'])} --> "
                      f"{fmt_ts(seg['end_s'])}\n{text}\n")
    out.write_text("\n".join(blocks), encoding="utf-8")
    return out


def export_vtt(session_id: int, out: Path, opts: ExportOptions) -> Path:
    d = load_session_data(session_id, opts)
    src = d.translation if opts.include_translation and d.translation \
        else d.segments
    blocks = ["WEBVTT", ""]
    for seg in src:
        text = seg["text"]
        if opts.include_speakers and seg.get("speaker"):
            text = f"<v {seg['speaker']}>{text}"
        blocks.append(f"{fmt_ts(seg['start_s'], vtt=True)} --> "
                      f"{fmt_ts(seg['end_s'], vtt=True)}\n{text}\n")
    out.write_text("\n".join(blocks), encoding="utf-8")
    return out


def export_json(session_id: int, out: Path, opts: ExportOptions) -> Path:
    d = load_session_data(session_id, opts)
    out.write_text(json.dumps({
        "session": d.name, "date": d.date, "duration_s": d.duration_s,
        "languages": d.languages, "provider": d.provider,
        "segments": d.segments,
        "translation": d.translation or None,
        "summary": d.summary or None,
        "exported_by": APP_NAME,
        "exported_at": dt.datetime.now().isoformat(timespec="seconds"),
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    return out


def export_markdown(session_id: int, out: Path,
                    opts: ExportOptions) -> Path:
    d = load_session_data(session_id, opts)
    lines = [f"# {opts.title or d.name}", "",
             f"*{d.date} — duration {fmt_ts(d.duration_s)[:8]} — "
             f"languages: {d.languages or 'n/a'}*", ""]
    if d.summary:
        lines += ["## Summary (AI-generated)", _summary_text(d.summary),
                  ""]
    lines.append("## Transcript")
    for seg in d.segments:
        lines.append("- " + _line(seg, opts))
    if d.translation:
        lines += ["", f"## Translation ({opts.include_translation})"]
        for seg in d.translation:
            lines.append("- " + _line(seg, opts))
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def export_html(session_id: int, out: Path, opts: ExportOptions) -> Path:
    d = load_session_data(session_id, opts)
    esc = html_mod.escape
    rows = []
    for seg in d.segments:
        rows.append(
            "<tr>"
            + (f"<td class=ts>{fmt_ts(seg['start_s'])[:8]}</td>"
               if opts.include_timestamps else "")
            + (f"<td class=spk>{esc(seg.get('speaker', ''))}</td>"
               if opts.include_speakers else "")
            + f"<td>{esc(seg['text'])}</td></tr>")
    wm = (f"<div class=wm>{esc(opts.watermark)}</div>"
          if opts.watermark else "")
    summary_html = (f"<h2>Summary (AI-generated)</h2><pre>"
                    f"{esc(_summary_text(d.summary))}</pre>"
                    if d.summary else "")
    out.write_text(f"""<!doctype html><html><head><meta charset=utf-8>
<title>{esc(d.name)}</title><style>
body{{font-family:Segoe UI,Arial,sans-serif;margin:2em;color:#222}}
table{{border-collapse:collapse;width:100%}}
td{{padding:4px 8px;vertical-align:top;border-bottom:1px solid #eee}}
.ts{{color:#888;white-space:nowrap}}.spk{{font-weight:600;
white-space:nowrap}}
.wm{{position:fixed;top:40%;left:20%;font-size:72px;color:#00000014;
transform:rotate(-30deg);pointer-events:none}}
</style></head><body>{wm}<h1>{esc(d.name)}</h1>
<p>{esc(d.date)} — {fmt_ts(d.duration_s)[:8]} — {esc(d.languages)}</p>
{summary_html}<h2>Transcript</h2><table>{''.join(rows)}</table>
<footer><p>{esc(opts.footer)}</p></footer></body></html>""",
                   encoding="utf-8")
    return out


# ------------------------------------------------------------------ DOCX

def export_docx(session_id: int, out: Path, opts: ExportOptions) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor
    d = load_session_data(session_id, opts)
    doc = Document()
    doc.add_heading(opts.title or d.name, level=0)
    meta = doc.add_paragraph(f"{d.date} — duration "
                             f"{fmt_ts(d.duration_s)[:8]} — languages: "
                             f"{d.languages or 'n/a'}")
    meta.runs[0].font.size = Pt(9)
    if opts.watermark:
        p = doc.add_paragraph(opts.watermark)
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        p.runs[0].font.bold = True
    if d.summary:
        doc.add_heading("Summary (AI-generated)", level=1)
        for para in _summary_text(d.summary).split("\n"):
            if para.strip():
                doc.add_paragraph(para)
    doc.add_heading("Transcript", level=1)
    for seg in d.segments:
        p = doc.add_paragraph()
        if opts.include_timestamps:
            r = p.add_run(f"[{fmt_ts(seg['start_s'])[:8]}] ")
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            r.font.size = Pt(8)
        if opts.include_speakers and seg.get("speaker"):
            r = p.add_run(f"{seg['speaker']}: ")
            r.bold = True
        p.add_run(seg["text"])
    if d.translation:
        doc.add_heading(f"Translation ({opts.include_translation})",
                        level=1)
        for seg in d.translation:
            p = doc.add_paragraph()
            if opts.include_speakers and seg.get("speaker"):
                p.add_run(f"{seg['speaker']}: ").bold = True
            p.add_run(seg["text"])
    if opts.footer:
        doc.add_paragraph(opts.footer)
    doc.save(str(out))
    return out


# ------------------------------------------------------------------ XLSX

def export_xlsx(session_id: int, out: Path, opts: ExportOptions) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
    d = load_session_data(session_id, opts)
    wb = Workbook()
    ws = wb.active
    ws.title = "Transcript"
    headers = ["#", "Start", "End", "Speaker", "Language", "Text"]
    ws.append(headers)
    fill = PatternFill("solid", fgColor="1F4E79")
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = fill
    for i, seg in enumerate(d.segments, 1):
        ws.append([i, fmt_ts(seg["start_s"])[:8], fmt_ts(seg["end_s"])[:8],
                   seg.get("speaker", ""), seg.get("language", ""),
                   seg["text"]])
    widths = [5, 10, 10, 18, 10, 100]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    for row in ws.iter_rows(min_row=2):
        row[5].alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    if d.summary:
        ws2 = wb.create_sheet("Summary")
        ws2.append(["Field", "Content (AI-generated)"])
        for c in ws2[1]:
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = fill
        for k, v in d.summary.items():
            ws2.append([k, _flatten(v)])
        ws2.column_dimensions["A"].width = 24
        ws2.column_dimensions["B"].width = 110
        for row in ws2.iter_rows(min_row=2):
            row[1].alignment = Alignment(wrap_text=True, vertical="top")
        ws2.freeze_panes = "A2"
    wb.save(str(out))
    return out


# ------------------------------------------------------------------- PDF

def export_pdf(session_id: int, out: Path, opts: ExportOptions) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer)
    d = load_session_data(session_id, opts)
    styles = getSampleStyleSheet()
    seg_style = ParagraphStyle("seg", parent=styles["Normal"],
                               fontSize=9.5, leading=13)
    story = [Paragraph(html_mod.escape(opts.title or d.name),
                       styles["Title"]),
             Paragraph(html_mod.escape(
                 f"{d.date} — {fmt_ts(d.duration_s)[:8]} — "
                 f"{d.languages or 'n/a'}"), styles["Normal"]),
             Spacer(1, 6 * mm)]
    if d.summary:
        story.append(Paragraph("Summary (AI-generated)",
                               styles["Heading1"]))
        for para in _summary_text(d.summary).split("\n"):
            if para.strip():
                story.append(Paragraph(html_mod.escape(para), seg_style))
        story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("Transcript", styles["Heading1"]))
    for seg in d.segments:
        prefix = ""
        if opts.include_timestamps:
            prefix += f"<font color=grey size=7>" \
                      f"[{fmt_ts(seg['start_s'])[:8]}]</font> "
        if opts.include_speakers and seg.get("speaker"):
            prefix += f"<b>{html_mod.escape(seg['speaker'])}:</b> "
        story.append(Paragraph(prefix + html_mod.escape(seg["text"]),
                               seg_style))

    def _decorate(canvas, doc_):
        canvas.saveState()
        if opts.watermark:
            canvas.setFont("Helvetica-Bold", 60)
            canvas.setFillColorRGB(0.9, 0.9, 0.9)
            canvas.translate(A4[0] / 2, A4[1] / 2)
            canvas.rotate(35)
            canvas.drawCentredString(0, 0, opts.watermark)
        canvas.setFillColorRGB(0.5, 0.5, 0.5)
        canvas.setFont("Helvetica", 7)
        canvas.drawString(15 * mm, 10 * mm, opts.footer or
                          f"Generated by {APP_NAME}")
        canvas.drawRightString(A4[0] - 15 * mm, 10 * mm,
                               f"Page {doc_.page}")
        canvas.restoreState()

    SimpleDocTemplate(str(out), pagesize=A4).build(
        story, onFirstPage=_decorate, onLaterPages=_decorate)
    return out


# ---------------------------------------------------------------- helpers

def _flatten(v) -> str:
    if isinstance(v, list):
        return "\n".join(_flatten(x) for x in v)
    if isinstance(v, dict):
        return "; ".join(f"{k}: {x}" for k, x in v.items() if x)
    return str(v)


def _summary_text(summary: dict) -> str:
    order = ["executive_summary", "detailed_summary", "topics",
             "decisions", "action_items", "questions_open", "risks",
             "quotes", "sentiment", "next_steps"]
    lines = []
    for key in order + [k for k in summary if k not in order]:
        if key not in summary or not summary[key]:
            continue
        title = key.replace("_", " ").title()
        lines.append(f"{title}:")
        lines.append(_flatten(summary[key]))
        lines.append("")
    return "\n".join(lines).strip()


EXPORTERS = {
    "txt": export_txt, "srt": export_srt, "vtt": export_vtt,
    "json": export_json, "md": export_markdown, "html": export_html,
    "docx": export_docx, "xlsx": export_xlsx, "pdf": export_pdf,
}


def export_session(session_id: int, fmt: str, out_path: Path,
                   opts: ExportOptions | None = None) -> Path:
    fmt = fmt.lower().lstrip(".")
    if fmt == "doc":
        fmt = "docx"
    if fmt not in EXPORTERS:
        raise ValueError(f"Unsupported export format: {fmt}")
    out_path = out_path.with_suffix(f".{fmt}")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    return EXPORTERS[fmt](session_id, out_path, opts or ExportOptions())
